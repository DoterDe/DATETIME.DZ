import os
import json
from docx import Document
import openpyxl
import requests
import time
from openpyxl import load_workbook

# wb=openpyxl.Workbook()
# sheet=wb.active #получение актив страницы
# for i in range(1,11):
#     sheet.cell(row=1, column=i).value = i
#     sheet.cell(row=2, column=i).value = chr(96+i)

# wb.save("exel_file.xlsx")

# file=open('exel_file.xlsx')

# file.close()

# with open('exel_file.xlsx') as file:
#     data = file.readline()
#     print(data)

# file= load_workbook('exel_file.xlsx')
# active =  file.active
# new=openpyxl.Workbook()
# new_active = new.active
# for i in range(1,11):
#     new_active.cell(row=i, column=1).value = active.cell(row=1, column=i).value
#     new_active.cell(row=i, column=2).value = active.cell(row=2, column=i).value
# new.save('new-exel.xlsx')

# response=requests.get("")
# print(response.json())
# with open('dta.json', mode='w') as file:
#     file.write(response.text)




doc=Document()
doc.add_heading('Document Title', 0)
doc.add_paragraph("Hello")
doc.save('word.docx')


# print(os.chdir(r'C:\Users\chila\Desktop'))

# print(os.path.exists(r'C:\Users\chila\Desktop\test'))
# print('              -------------------------------',os.getcwd())
# print(os.path.isfile(r'C:\Users\chila\Desktop\test'))
# os.makedirs('Folder')
# time.sleep(3)
# os.chmod('classwork.py', 777)
# os.rmdir('Folder')